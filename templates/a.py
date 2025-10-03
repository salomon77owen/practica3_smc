from docx import Document
import datetime

# Crear documento
doc = Document()
doc.add_heading("Guía para subir un proyecto Flask a GitHub y Render", 0)

# Fecha
fecha = datetime.datetime.now().strftime("%d/%m/%Y")
doc.add_paragraph(f"Fecha: {fecha}")

# Sección 1: Subir a GitHub
doc.add_heading("1. Subir el proyecto a GitHub", level=1)

doc.add_paragraph("Asegúrate de tener una cuenta en GitHub y Git instalado. Luego sigue estos pasos desde tu terminal:")

steps_github = [
    "1.1 Abre la terminal y navega a la carpeta de tu proyecto:",
    "    cd ruta/a/tu/proyecto",
    "1.2 Inicializa Git (si no lo hiciste antes):",
    "    git init",
    "1.3 Agrega todos los archivos:",
    "    git add .",
    "1.4 Crea un commit con un mensaje:",
    "    git commit -m \"Primer commit\"",
    "1.5 Crea tu repositorio en GitHub (desde la web) y copia la URL.",
    "1.6 Conecta tu proyecto local al repositorio remoto:",
    "    git remote add origin https://github.com/usuario/repositorio.git",
    "1.7 Cambia el nombre de la rama principal (opcional):",
    "    git branch -M main",
    "1.8 Sube el proyecto a GitHub:",
    "    git push -u origin main"
]
for step in steps_github:
    doc.add_paragraph(step, style='ListBullet')

# Sección 2: Preparar proyecto para Render
doc.add_heading("2. Preparar el proyecto para Render", level=1)
doc.add_paragraph("Render necesita algunos archivos específicos:")

doc.add_paragraph("- requirements.txt:")
doc.add_paragraph("  Ejecuta este comando en tu terminal para generarlo:")
doc.add_paragraph("      pip freeze > requirements.txt")

doc.add_paragraph("- Procfile (sin extensión):")
doc.add_paragraph("  Crea un archivo con este contenido:")
doc.add_paragraph("      web: gunicorn app:app")

doc.add_paragraph("- Asegúrate de tener esta estructura de proyecto:")
doc.add_paragraph("""
/mi-proyecto/
│
├── app.py
├── requirements.txt
├── Procfile
├── templates/
│   └── index.html
├── static/
│   ├── css/
│   │   └── estilo.css
│   └── img/
│       └── imagen.jpg
""")

# Sección 3: Subir a Render
doc.add_heading("3. Subir el proyecto a Render", level=1)

steps_render = [
    "3.1 Ve a https://render.com y crea una cuenta o inicia sesión.",
    "3.2 Haz clic en 'New' → 'Web Service'.",
    "3.3 Conecta tu cuenta de GitHub y selecciona el repositorio de tu proyecto.",
    "3.4 Completa la configuración del servicio:",
    "     - Runtime: Python",
    "     - Build Command: pip install -r requirements.txt",
    "     - Start Command: gunicorn app:app",
    "     - Branch: main (u otra que uses)",
    "3.5 Haz clic en 'Create Web Service'.",
    "3.6 Espera que Render construya y despliegue tu app.",
    "3.7 Render te mostrará una URL como https://tu-proyecto.onrender.com"
]
for step in steps_render:
    doc.add_paragraph(step, style='ListBullet')

doc.add_paragraph("¡Listo! Tu proyecto Flask ya estará en línea.")

# Guardar documento
doc.save("Guia_Subir_Proyecto_Flask_GitHub_Render.docx")
